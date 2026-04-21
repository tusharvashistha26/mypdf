from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles

from pdf2docx import Converter
from PIL import Image
from pypdf import PdfReader, PdfWriter
import platform
from docx import Document
from weasyprint import HTML

import subprocess
import os
import uuid
import shutil
import threading
import multiprocessing

workers = max(1, multiprocessing.cpu_count() // 2)

app = FastAPI()

templates = Jinja2Templates(directory="templates")

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

app.mount("/static", StaticFiles(directory="static"), name="static")

jobs = {}
lock = threading.Lock()

# =============================
# HOME
# =============================
@app.get("/")
def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


# =============================
# FAST FILE SAVE
# =============================
def save_file(file: UploadFile):
    path = f"{UPLOAD_DIR}/{uuid.uuid4()}_{file.filename}"
    with open(path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    return path


# =============================
# SAFE JOB RUNNER
# =============================
def run_job(job_id, func, *args):
    try:
        output_path = func(*args)

        if not os.path.exists(output_path):
            raise Exception("Output file missing")

        with lock:
            jobs[job_id] = {
                "status": "completed",
                "file": output_path,
                "download_url": f"/download/{job_id}"
            }

    except Exception as e:
        print("ERROR:", str(e))
        with lock:
            jobs[job_id] = {
                "status": "failed",
                "message": str(e)
            }


# =============================
# PROCESS FUNCTIONS
# =============================

def process_images(paths):
    images = []
    for p in paths:
        img = Image.open(p)
        img = img.convert("RGB")
        img.thumbnail((2000, 2000))  # 🔥 reduce size
        images.append(img)

    output = f"{OUTPUT_DIR}/{uuid.uuid4()}.pdf"
    images[0].save(output, save_all=True, append_images=images[1:])

    return output


def process_compress(path, level):
    output = f"{OUTPUT_DIR}/{uuid.uuid4()}.pdf"

    result = subprocess.run([
        "gs",
        "-sDEVICE=pdfwrite",
        "-dCompatibilityLevel=1.4",
        f"-dPDFSETTINGS={level}",
        "-dNOPAUSE",
        "-dQUIET",
        "-dBATCH",
        f"-sOutputFile={output}",
        path
    ], capture_output=True, text=True)

    if result.returncode != 0:
        raise Exception(result.stderr or "Compression failed")

    return output


def process_merge(paths):
    writer = PdfWriter()

    for p in paths:
        reader = PdfReader(p)
        for page in reader.pages:
            writer.add_page(page)

    output = f"{OUTPUT_DIR}/{uuid.uuid4()}.pdf"

    with open(output, "wb") as f:
        writer.write(f)

    return output


def process_split(path, page):
    reader = PdfReader(path)

    if page < 1 or page > len(reader.pages):
        raise Exception("Invalid page number")

    writer = PdfWriter()
    writer.add_page(reader.pages[page - 1])

    output = f"{OUTPUT_DIR}/{uuid.uuid4()}.pdf"

    with open(output, "wb") as f:
        writer.write(f)

    return output


# =============================
# 🚀 PURE PYTHON WORD → PDF
# =============================

def process_word(path):
    output = f"{OUTPUT_DIR}/{uuid.uuid4()}.pdf"

    # 🪟 WINDOWS → fallback (simple text PDF)
    if platform.system() == "Windows":
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from docx import Document

        doc = Document(path)
        c = canvas.Canvas(output, pagesize=letter)

        y = 750
        for para in doc.paragraphs:
            c.drawString(50, y, para.text[:90])
            y -= 20
            if y < 50:
                c.showPage()
                y = 750

        c.save()
        return output

    # 🐧 LINUX (Render) → WeasyPrint (FULL QUALITY)
    else:
        from docx import Document
        from weasyprint import HTML

        doc = Document(path)

        html = ""
        for p in doc.paragraphs:
            html += f"<p>{p.text}</p>"

        HTML(string=html).write_pdf(output)
        return output


# =============================
# PDF → WORD
# =============================
def process_pdf_to_word(path):
    output = f"{OUTPUT_DIR}/{uuid.uuid4()}.docx"

    cv = Converter(path)
    cv.convert(output)
    cv.close()

    return output


# =============================
# ENQUEUE APIs
# =============================
@app.post("/enqueue/images-to-pdf")
async def enqueue_images(background_tasks: BackgroundTasks, files: list[UploadFile] = File(...)):
    job_id = str(uuid.uuid4())
    paths = [save_file(file) for file in files]

    jobs[job_id] = {"status": "processing"}
    background_tasks.add_task(run_job, job_id, process_images, paths)

    return {"job_id": job_id}


@app.post("/enqueue/compress-pdf")
async def enqueue_compress(background_tasks: BackgroundTasks, file: UploadFile = File(...), level: str = Form(...)):
    job_id = str(uuid.uuid4())
    path = save_file(file)

    jobs[job_id] = {"status": "processing"}
    background_tasks.add_task(run_job, job_id, process_compress, path, level)

    return {"job_id": job_id}


@app.post("/enqueue/merge-pdf")
async def enqueue_merge(background_tasks: BackgroundTasks, files: list[UploadFile] = File(...)):
    job_id = str(uuid.uuid4())
    paths = [save_file(file) for file in files]

    jobs[job_id] = {"status": "processing"}
    background_tasks.add_task(run_job, job_id, process_merge, paths)

    return {"job_id": job_id}


@app.post("/enqueue/split-pdf")
async def enqueue_split(background_tasks: BackgroundTasks, file: UploadFile = File(...), page: int = Form(...)):
    job_id = str(uuid.uuid4())
    path = save_file(file)

    jobs[job_id] = {"status": "processing"}
    background_tasks.add_task(run_job, job_id, process_split, path, page)

    return {"job_id": job_id}


@app.post("/enqueue/word-to-pdf")
async def enqueue_word(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    job_id = str(uuid.uuid4())
    path = save_file(file)

    jobs[job_id] = {"status": "processing"}
    background_tasks.add_task(run_job, job_id, process_word, path)

    return {"job_id": job_id}


@app.post("/enqueue/pdf-to-word")
async def enqueue_pdf(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    job_id = str(uuid.uuid4())
    path = save_file(file)

    jobs[job_id] = {"status": "processing"}
    background_tasks.add_task(run_job, job_id, process_pdf_to_word, path)

    return {"job_id": job_id}


# =============================
# STATUS
# =============================
@app.get("/job-status/{job_id}")
def job_status(job_id: str):
    return JSONResponse(jobs.get(job_id, {"status": "not_found"}))


# =============================
# DOWNLOAD
# =============================
@app.get("/download/{job_id}")
def download(job_id: str):
    job = jobs.get(job_id)

    if not job or job["status"] != "completed":
        raise HTTPException(400, "Not ready")

    return FileResponse(job["file"], filename=os.path.basename(job["file"]))


# =============================
# GLOBAL ERROR FIX (NO HTML ERRORS)
# =============================
@app.exception_handler(Exception)
def global_exception(request, exc):
    return JSONResponse(
        status_code=500,
        content={"status": "error", "message": str(exc)}
    )